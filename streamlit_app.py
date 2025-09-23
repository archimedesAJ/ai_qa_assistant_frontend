# import time
# import streamlit as st
# import pandas as pd
# import base64
# from io import BytesIO
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# import requests

# API_BASE = "http://127.0.0.1:8000/api"

# # --- Logo in Sidebar ---
# st.sidebar.image("images/purplewave-logo.jpg", use_container_width=True)

# # --- App Title ---
# st.sidebar.title("PW QA Assistant Platform")

# # --- Project Selection ---
# # --- Project Selection ---
# st.sidebar.subheader("Select Your Team")

# try:
#     response = requests.get(f"{API_BASE}/teams/")
#     if response.status_code == 200:
#         teams = response.json()
#         team_options = {team["name"]: team["id"] for team in teams}
#         project_name = st.sidebar.selectbox("Available Teams", list(team_options.keys()))
#         selected_team_id = team_options[project_name]
#     else:
#         st.sidebar.error("⚠️ Could not fetch teams")
#         project_name = None
#         selected_team_id = None
# except Exception as e:
#     st.sidebar.error(f"⚠️ API Error: {e}")
#     project_name = None
#     selected_team_id = None

# # --- Tabs ---
# tab = st.sidebar.radio("Navigation", ["Test Cases", "Test Plan"])

# # --- Data Storage ---
# if "test_cases" not in st.session_state:
#     st.session_state.test_cases = []

# if "file_uploader_key" not in st.session_state:
#     st.session_state.file_uploader_key = 0

# if "test_plan" not in st.session_state:
#     st.session_state.test_plan = ""


# # --- Test Cases Tab ---
# if tab == "Test Cases":
#     st.title("Test Case Generator")
#     st.write("Generate test cases from user story")
#     user_story = st.text_input("Story Title *", key="story_title")
#     description = st.text_area("Description *", key="description")
#     acceptance_criteria = st.text_area("Acceptance Criteria *", key="acceptance_criteria")

#     error_placeholder = st.empty()  # placeholder for error

#     if st.button("Generate Test Cases"):
#         if not user_story.strip() or not acceptance_criteria.strip():
#             error_placeholder.error("⚠️ Please fill in both *User Story* and *Acceptance Criteria* before generating test cases.")
#             time.sleep(3)  # wait 3 seconds
#             error_placeholder.empty()  # clear the alert
#             st.rerun()  # refresh after alert disappears
#         else:
#             # Placeholder Gherkin generator
#             # test_case = {
#             #     "title": "Verify login functionality",
#             #     "description": f"Story: {user_story}\nCriteria: {acceptance_criteria}",
#             #     "gherkin": """Feature: Login
#             #     Scenario: Valid user login
#             #         Given user is on login page
#             #         When user enters valid credentials
#             #         Then dashboard is displayed"""
#             # }
#             # st.session_state.test_cases.append(test_case)

#             # --- Generate multiple test cases ---
#             test_cases = [
#                 {
#                     "title": "Verify login with valid credentials",
#                     "description": f"Story: {user_story}\nCriteria: {acceptance_criteria}",
#                     "gherkin": """Feature: Login
#                 Scenario: Valid user login
#                     Given user is on login page
#                     When user enters valid credentials
#                     Then dashboard is displayed"""
#                 },
#                 {
#                     "title": "Verify login with invalid password",
#                     "description": f"Story: {user_story}\nCriteria: {acceptance_criteria}",
#                     "gherkin": """Feature: Login
#                 Scenario: Invalid password
#                     Given user is on login page
#                     When user enters valid username and invalid password
#                     Then error message is displayed"""
#                 },
#                 {
#                     "title": "Verify login with empty fields",
#                     "description": f"Story: {user_story}\nCriteria: {acceptance_criteria}",
#                     "gherkin": """Feature: Login
#                 Scenario: Empty fields
#                     Given user is on login page
#                     When user clicks login without entering credentials
#                     Then error message is displayed"""
#                 }
#             ]

#             st.session_state.test_cases.extend(test_cases)


#     if st.session_state.test_cases:
#         st.subheader("Generated Test Cases")
#         for i, tc in enumerate(st.session_state.test_cases):
#             with st.expander(tc["title"]):
#                 st.write(tc["description"])
#                 st.code(tc["gherkin"], language="gherkin")
#                 if st.button(f"Delete {i}"):
#                     st.session_state.test_cases.pop(i)
#                     st.rerun()

#         # --- Row with Export, Send to Jira, Reset ---
#         col1, col2, col3 = st.columns([6, 1, 1])

#         # Export button
#         with col1:
#             if st.button("⬇️ Export All"):
#                 df = pd.DataFrame(st.session_state.test_cases)
#                 csv = df.to_csv(index=False).encode('utf-8')
#                 st.download_button(
#                     "Download CSV",
#                     data=csv,
#                     file_name="test_cases.csv",
#                     mime="text/csv"
#                 )

#         jira_placeholder = st.empty()  # placeholder for Jira alert

#         # Send to Jira button
#         with col2:
#             if st.button("📨", help="Send to Jira"):
#                 jira_placeholder.success("✅ Sent to Jira")
#                 time.sleep(3)
#                 jira_placeholder.empty()
#                 st.rerun()

#         # Reset Test Case Generator button
#         with col3:
#             if st.button("🔄", help="Reset Test Case Generator"):
#                 st.session_state.clear()  # clear everything
#                 st.rerun() 


# # --- Test Plan Tab ---
# if tab == "Test Plan":
#     st.title("Test Plan Generator")

#     uploaded_files = st.file_uploader(
#         "Upload Documents [BRD / SRS / FSD]",
#         type=["docx", "pdf"],   # restrict file types
#         accept_multiple_files=True,
#         key=f"file_uploader_{st.session_state.file_uploader_key}"
#     )

#     if uploaded_files:
#         st.subheader("Uploaded Files")
#         content = ""
#         for file in uploaded_files:
#             content += file.read().decode("utf-8", errors="ignore") + "\n"
#         st.text_area("Combined Content", content, height=200)

#         if st.button("Generate Test Plan"):
#             st.session_state.test_plan = f"Test Plan for {project}\n\n{content[:500]}..."

#     if st.session_state.test_plan:
#         st.subheader("Generated Test Plan")
#         st.text_area("Test Plan", st.session_state.test_plan, height=300)

#         # --- Right-aligned Reset Button (icon only) ---
#         col1, col2 = st.columns([8, 1])
#         with col2:
#             if st.button("🔄", help="Reset Test Plan"):
#                 st.session_state.test_plan = ""
#                 st.session_state.file_uploader_key += 1
#                 st.rerun()

#         # Export DOCX
#         if st.button("Export as DOCX"):
#             doc = Document()
#             doc.add_heading(f"Test Plan - {project}", 0)
#             doc.add_paragraph(st.session_state.test_plan)
#             buffer = BytesIO()
#             doc.save(buffer)
#             buffer.seek(0)
#             st.download_button("Download DOCX", buffer, file_name="test_plan.docx")

#         # Export PDF
#         if st.button("Export as PDF"):
#             buffer = BytesIO()
#             c = canvas.Canvas(buffer, pagesize=letter)
#             text_object = c.beginText(40, 750)
#             for line in st.session_state.test_plan.split("\n"):
#                 text_object.textLine(line)
#             c.drawText(text_object)
#             c.showPage()
#             c.save()
#             buffer.seek(0)
#             st.download_button("Download PDF", buffer, file_name="test_plan.pdf")

# # --- Footer ---
# st.markdown(
#     """
#     <hr style="margin-top: 50px; margin-bottom: 10px;">
#     <div style="text-align: center; color: grey; font-size: 14px;">
#         © 2025 QA Assistant Platform. Built with ❤️ by Team 3
#     </div>
#     """,
#     unsafe_allow_html=True
# )

import json
import time
import streamlit as st
import pandas as pd
import base64
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import requests

API_BASE = "http://127.0.0.1:8000/api"

# --- Logo in Sidebar ---
st.sidebar.image("images/purplewave-logo.jpg", use_container_width=True)

# --- App Title ---
st.sidebar.title("PW QA Assistant Platform")

# --- Project Selection ---
st.sidebar.subheader("Select Your Team")

try:
    response = requests.get(f"{API_BASE}/teams/")
    if response.status_code == 200:
        teams = response.json()
        team_options = {team["name"]: team["id"] for team in teams}
        project_name = st.sidebar.selectbox("Available Teams", list(team_options.keys()))
        selected_team_id = team_options[project_name]
    else:
        st.sidebar.error("⚠️ Could not fetch teams")
        project_name = None
        selected_team_id = None
except Exception as e:
    st.sidebar.error(f"⚠️ API Error: {e}")
    project_name = None
    selected_team_id = None

# --- Tabs ---
tab = st.sidebar.radio("Navigation", ["Test Cases", "Test Plan", "PW QA Guide"])

# --- Data Storage ---
if "test_cases" not in st.session_state:
    st.session_state.test_cases = []

if "file_uploader_key" not in st.session_state:
    st.session_state.file_uploader_key = 0

if "test_plan" not in st.session_state:
    st.session_state.test_plan = ""

# --- Chat Interface Data Storage ---
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []

if "show_feedback" not in st.session_state:
    st.session_state.show_feedback = {}

# --- Test Cases Tab ---
if tab == "Test Cases":
    st.title("Test Case Generator")

    # Create two sub-tabs
    subtab_manual, subtab_story_key = st.tabs(["✍️ Manual Entry", "🔑 From Story Key"])
    
    # --------------------
    # Sub-tab 1: Manual Entry
    # --------------------
    with subtab_manual:
        st.subheader("Manual Entry")
        st.write("Fill in story details manually and generate test cases from the user story") 
        
        user_story = st.text_input("Story Title *", key="story_title")
        description = st.text_area("Description *", key="description")
        acceptance_criteria = st.text_area("Acceptance Criteria *", key="acceptance_criteria")

        error_placeholder = st.empty()  # placeholder for error

        if st.button("🚀Generate Test Cases", key="generate_manual"):
            if not user_story.strip() or not acceptance_criteria.strip():
                error_placeholder.error("⚠️ Please fill in both *User Story* and *Acceptance Criteria* before generating test cases.")
                time.sleep(3)  # wait 3 seconds
                error_placeholder.empty()  # clear the alert
                st.rerun()  # refresh after alert disappears
            else:
                # --- Progress bar setup ---
                progress_bar = st.progress(0)
                status_text = st.empty()

                payload = {
                    "user_story": user_story,
                    "acceptance_criteria": acceptance_criteria,
                    "feature_description": description,
                    "team_id": selected_team_id
                }

                try:
                    # Fake progress updates while waiting
                    messages = [
                        "🔍 Analyzing requirements...",
                        "🧩 Designing scenarios...",
                        "✍️ Writing Gherkin steps...",
                        "🛠️ Validating test coverage...",
                        "📦 Finalizing test cases..."
                    ]
                    steps = [10, 30, 50, 70, 90]

                    for pct, msg in zip(steps, messages):
                        progress_bar.progress(pct)
                        status_text.text(msg)
                        time.sleep(5)

                    # --- Actual API call ---
                    response = requests.post(
                        f"{API_BASE}/generate/testcases/prompt/",
                        json=payload,
                        timeout=120
                    )

                    if response.status_code == 200:
                        raw_text = response.json().get("test_cases", "")

                        # Clean up if response wrapped in ```json ```
                        if raw_text.startswith("```json"):
                            raw_text = raw_text.strip("`").replace("json", "", 1).strip()

                        parsed = {}
                        try:
                            parsed = json.loads(raw_text)
                        except Exception as e:
                            st.error(f"⚠️ Failed to parse response: {e}")

                        test_cases = parsed.get("test_cases", [])
                        st.session_state.test_cases.extend(test_cases)

                        progress_bar.progress(100)
                        status_text.text("✅ Test cases generated successfully!")

                    else:
                        st.error(f"⚠️ API Error: {response.status_code}")
                        status_text.text("❌ Failed to generate test cases")

                except Exception as e:
                    st.error(f"⚠️ Request Failed: {e}")
                    status_text.text("❌ Error during test case generation")

    with subtab_story_key:
        st.subheader("From Jira/Story Key")
        st.write("Enter a user story key (e.g. FP-1325) to fetch details automatically")
        
        story_key = st.text_input("Story Key *", placeholder="e.g. FP-1325", key="story_key_input")
        
        if story_key.strip():
            try:
                with st.spinner("Fetching story details..."):
                    response = requests.post(
                        f"{API_BASE}/autopopulate/userstory/",
                        json={"story_key": story_key.strip()},
                        timeout=60
                    )
                story_alert = st.empty()
                if response.status_code == 200:
                    story_data = response.json()
                    story_alert.success("⌛ Loading story details...")
                    time.sleep(2)
                    story_alert.empty()
                
                    # Pre-fill with API response
                    user_story = st.text_input(
                        "Story Title *", 
                        value=story_data.get("summary", ""), 
                        key="story_title_key"
                    )
                    description = st.text_area(
                        "Description *", 
                        value=story_data.get("description", ""), 
                        key="description_key"
                    )
                    acceptance_criteria = st.text_area(
                        "Acceptance Criteria *", 
                        value=story_data.get("acceptance_criteria", ""), 
                        key="acceptance_key"
                    )
                    col1, col2 = st.columns([1,1])

                    with col1:
                        if st.button("🚀Generate Test Cases", key="generate_from_key"):
                            if not user_story.strip() or not acceptance_criteria.strip():
                                error_placeholder.error("⚠️ Please fill in both *User Story* and *Acceptance Criteria* before generating test cases.")
                                time.sleep(3)  # wait 3 seconds
                                error_placeholder.empty()  # clear the alert
                                st.rerun()  # refresh after alert disappears
                            else:
                                # --- Progress bar setup ---
                                progress_bar = st.progress(0)
                                status_text = st.empty()

                                payload = {
                                    "user_story": user_story,
                                    "acceptance_criteria": acceptance_criteria,
                                    "feature_description": description,
                                    "team_id": selected_team_id
                                }

                                try:
                                    # Fake progress updates while waiting
                                    messages = [
                                        "🔍 Analyzing requirements...",
                                        "🧩 Designing scenarios...",
                                        "✍️ Writing Gherkin steps...",
                                        "🛠️ Validating test coverage...",
                                        "📦 Finalizing test cases..."
                                    ]
                                    steps = [10, 30, 50, 70, 90]

                                    for pct, msg in zip(steps, messages):
                                        progress_bar.progress(pct)
                                        status_text.text(msg)
                                        time.sleep(5)

                                    # --- Actual API call ---
                                    response = requests.post(
                                        f"{API_BASE}/generate/testcases/prompt/",
                                        json=payload,
                                        timeout=120
                                    )

                                    if response.status_code == 200:
                                        raw_text = response.json().get("test_cases", "")

                                        # Clean up if response wrapped in ```json ```
                                        if raw_text.startswith("```json"):
                                            raw_text = raw_text.strip("`").replace("json", "", 1).strip()

                                        parsed = {}
                                        try:
                                            parsed = json.loads(raw_text)
                                        except Exception as e:
                                            st.error(f"⚠️ Failed to parse response: {e}")

                                        test_cases = parsed.get("test_cases", [])
                                        st.session_state.test_cases.extend(test_cases)

                                        progress_bar.progress(100)
                                        status_text.text("✅ Test cases generated successfully!")

                                    else:
                                        st.error(f"⚠️ API Error: {response.status_code}")
                                        status_text.text("❌ Failed to generate test cases")

                                except Exception as e:
                                    st.error(f"⚠️ Request Failed: {e}")
                                    status_text.text("❌ Error during test case generation")
                    
                else:
                    st.warning(f"⚠️ Could not fetch story (status {response.status_code})")
            except Exception as e:
                st.error(f"⚠️ Failed to fetch story: {e}")

    if st.session_state.test_cases:
        st.subheader("Generated Test Cases")
        for i, tc in enumerate(st.session_state.test_cases):
            with st.expander(tc.get("title", f"Test Case {i+1}")):
                st.write(tc.get("description", ""))
                if "gherkin" in tc:
                    st.code(tc["gherkin"], language="gherkin")
                else:
                    st.json(tc)  # fallback to raw JSON
                if st.button(f"Delete {i}"):
                    st.session_state.test_cases.pop(i)
                    st.rerun()

        # --- Row with Export, Send to Jira, Reset ---
        col1, col2, col3 = st.columns([6, 1, 1])

        with col1:
            if st.button("⬇️ Export All"):
                df = pd.DataFrame(st.session_state.test_cases)
                csv = df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    "Download CSV",
                    data=csv,
                    file_name="test_cases.csv",
                    mime="text/csv"
                )

        jira_placeholder = st.empty()

        with col2:
            if st.button("📨", help="Send to Jira"):
                jira_placeholder.success("✅ Sent to Jira")
                time.sleep(3)
                jira_placeholder.empty()
                st.rerun()

        with col3:
            if st.button("🔄", help="Reset Test Case Generator"):
                st.session_state.test_cases = []
                st.rerun()

# --- Test Plan Tab ---
elif tab == "Test Plan":
    st.title("Test Plan Generator")

    # Create sub-tabs: From Docs or From Jira Confluence
    docs_tab, jira_tab = st.tabs(["📂 From Documents", "🌐 From Jira Confluence"])

    # ==========================
    # 1. From Documents (existing)
    # ==========================
    with docs_tab:
        uploaded_files = st.file_uploader(
            "Upload Documents [BRD / SRS / FSD]",
            type=["docx", "pdf", "txt"],
            accept_multiple_files=True,
            key=f"file_uploader_{st.session_state.file_uploader_key}"
        )

        if uploaded_files and selected_team_id:
            # Two sub-tabs: Files list & Preview
            files_tab, preview_tab = st.tabs(["📂 Uploaded Files", "👀 Preview"])

            with files_tab:
                st.subheader("Uploaded Files")
                for file in uploaded_files:
                    st.write(f"- {file.name}")

                # --- Generate Test Plan ---
                if st.button("Generate Test Plan"):
                    try:
                        files = [("documents", (f.name, f, f.type)) for f in uploaded_files]
                        data = {"team_id": selected_team_id}

                        with st.spinner("⏳ Generating test plan..."):
                            response = requests.post(
                                f"{API_BASE}/generate/testplan/document/",
                                files=files,
                                data=data,
                                timeout=180
                            )

                        if response.status_code == 200:
                            result = response.json()
                            raw_plan = result.get("test_plan")

                            # handle string-wrapped JSON
                            if isinstance(raw_plan, str):
                                try:
                                    raw_plan = raw_plan.strip("`").replace("json", "", 1).strip()
                                    test_plan = json.loads(raw_plan)
                                except Exception as e:
                                    st.error(f"⚠️ Failed to parse test plan JSON: {e}")
                                    st.text_area("Raw Output", raw_plan, height=300)
                                    test_plan = None
                            else:
                                test_plan = raw_plan

                            if test_plan:
                                st.session_state.test_plan = test_plan
                                st.success("✅ Test Plan generated successfully!")
                        else:
                            st.error(f"⚠️ API Error {response.status_code}")
                            st.json(response.json())

                    except Exception as e:
                        st.error(f"⚠️ Request Failed: {e}")

            with preview_tab:
                st.subheader("📄 Document Preview")
                selected_file = st.selectbox(
                    "Choose a file to preview",
                    [file.name for file in uploaded_files]
                )

                if selected_file:
                    file_obj = next((f for f in uploaded_files if f.name == selected_file), None)
                    if file_obj:
                        file_type = file_obj.type
                        file_text = ""

                        if file_type == "application/pdf":
                            from PyPDF2 import PdfReader
                            pdf_reader = PdfReader(file_obj)
                            file_text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                            from docx import Document
                            doc = Document(file_obj)
                            file_text = "\n".join([para.text for para in doc.paragraphs])
                        else:  # txt or fallback
                            file_text = file_obj.read().decode("utf-8", errors="ignore")

                        st.markdown(f"### {selected_file}")
                        st.text_area("Content", file_text, height=400)


    # ==========================
    # 2. From Jira Confluence (new)
    # ==========================
    with jira_tab:
        st.subheader("🌐 Generate from Jira Confluence Link")

        confluence_url = st.text_input("Enter Jira Confluence page link", key="confluence_url")

        if confluence_url and selected_team_id:
            if st.button("Generate Test Plan", key="gen_from_confluence"):
                try:
                    data = {"team_id": selected_team_id, "confluence_url": confluence_url}

                    with st.spinner("⏳ Reading Confluence & generating test plan..."):
                        response = requests.post(
                            f"{API_BASE}/generate/testplan/confluence_url/",
                            json=data,
                            timeout=180
                        )

                    if response.status_code == 200:
                        result = response.json()
                        raw_plan = result.get("test_plan")

                        if isinstance(raw_plan, str):
                            try:
                                raw_plan = raw_plan.strip("`").replace("json", "", 1).strip()
                                test_plan = json.loads(raw_plan)
                            except Exception as e:
                                st.error(f"⚠️ Failed to parse test plan JSON: {e}")
                                st.text_area("Raw Output", raw_plan, height=300)
                                test_plan = None
                        else:
                            test_plan = raw_plan

                        if test_plan:
                            st.session_state.test_plan = test_plan
                            st.success("✅ Test Plan generated successfully from Confluence!")
                    else:
                        st.error(f"⚠️ API Error {response.status_code}")
                        st.json(response.json())

                except Exception as e:
                    st.error(f"⚠️ Request Failed: {e}")
    
    # --- Display Test Plan if available ---
    if st.session_state.test_plan:
        test_plan = st.session_state.test_plan

        st.subheader(f"🧩 Feature: {test_plan.get('feature', 'N/A')}")
        st.write("### 🎯 Objectives")
        st.write("\n".join([f"- {obj}" for obj in test_plan.get("objectives", [])]))

        st.write("### 📌 Scope")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**In Scope**")
            st.write("\n".join([f"- {item}" for item in test_plan.get("scope", {}).get("in_scope", [])]))
        with col2:
            st.markdown("**Out of Scope**")
            st.write("\n".join([f"- {item}" for item in test_plan.get("scope", {}).get("out_of_scope", [])]))

        st.write("### 🛠 Approach")
        st.write(test_plan.get("approach", ""))

        st.write("### ✅ Test Types")
        st.write(", ".join(test_plan.get("test_types", [])))

        st.write("### 🌍 Environments")
        st.write(", ".join(test_plan.get("environments", [])))

        st.write("### 📂 Data & Tools")
        st.markdown("**Test Data**")
        st.write("\n".join([f"- {d}" for d in test_plan.get("data_and_tools", {}).get("test_data", [])]))
        st.markdown("**Tools**")
        st.write("\n".join([f"- {t}" for t in test_plan.get("data_and_tools", {}).get("tools", [])]))

        st.write("### 👥 Roles & Responsibilities")
        st.write("\n".join([f"- {role}" for role in test_plan.get("roles_and_responsibilities", [])]))

        st.write("### ⚠️ Risks & Mitigations")
        for r in test_plan.get("risks_and_mitigations", []):
            st.markdown(f"- **Risk**: {r['risk']}  \n  **Mitigation**: {r['mitigation']}")

        st.write("### 🚪 Entry Criteria")
        st.write("\n".join([f"- {c}" for c in test_plan.get("entry_criteria", [])]))

        st.write("### ✅ Exit Criteria")
        st.write("\n".join([f"- {c}" for c in test_plan.get("exit_criteria", [])]))

        # --- Export & Reset Section ---
        st.markdown("---")
        col1, col2, col3 = st.columns([2, 2, 1])

        with col1:
            # Export DOCX
            doc = Document()
            doc.add_heading(f"Test Plan - {test_plan.get('feature', 'N/A')}", 0)

            doc.add_heading("Objectives", level=1)
            for obj in test_plan.get("objectives", []):
                doc.add_paragraph(obj)

            doc.add_heading("Scope", level=1)
            doc.add_paragraph("In Scope:")
            for item in test_plan.get("scope", {}).get("in_scope", []):
                doc.add_paragraph(f"- {item}")
            doc.add_paragraph("Out of Scope:")
            for item in test_plan.get("scope", {}).get("out_of_scope", []):
                doc.add_paragraph(f"- {item}")

            doc.add_heading("Approach", level=1)
            doc.add_paragraph(test_plan.get("approach", ""))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("⬇️ Download DOCX", buffer, file_name="test_plan.docx")

        with col2:
            # Export PDF
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            text_object = c.beginText(40, 750)
            text_object.setFont("Helvetica", 10)

            # Dump test plan as text
            for section, value in test_plan.items():
                text_object.textLine(f"{section.upper()}:")
                if isinstance(value, list):
                    for v in value:
                        text_object.textLine(f"- {v}")
                elif isinstance(value, dict):
                    for k, v in value.items():
                        text_object.textLine(f"{k}: {v}")
                else:
                    text_object.textLine(str(value))
                text_object.textLine("")

            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button("⬇️ Download PDF", buffer, file_name="test_plan.pdf")

        with col3:
            # Reset everything
            if st.button("🔄 Reset", help="Clear generated plan & uploads"):
                st.session_state.test_plan = ""
                st.session_state.file_uploader_key += 1  # refresh uploader
                st.rerun()

# --- PW QA Guide Tab (NEW CHAT INTERFACE) ---
elif tab == "PW QA Guide":
    st.title("🤖 PW QA Assistant Guide")
    st.write(f"Get instant help with QA processes, testing procedures, and {project_name} specific guidance.")
    
    # --- Chat Helper Functions ---
    def get_ai_response(question, team_id):
        """Get AI response from backend API"""
        try:
            payload = {
                "question": question,
                "project_team_id": team_id,
                "user_context": {
                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "team_name": project_name
                }
            }
            
            response = requests.post(
                f"{API_BASE}/qa-assistant/",
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"API Error: {response.status_code}"}
                
        except Exception as e:
            return {"error": f"Request failed: {str(e)}"}
    
    def submit_feedback(query_id, rating):
        """Submit user feedback for AI response"""
        try:
            payload = {
                "query_id": query_id,
                "rating": rating
            }
            
            response = requests.post(
                f"{API_BASE}/qa-feedback/",
                json=payload,
                timeout=10
            )
            
            if response.status_code == 200:
                st.success("✅ Feedback submitted!")
            else:
                st.error("⚠️ Failed to submit feedback")
                
        except Exception as e:
            st.error(f"⚠️ Feedback error: {e}")
    
    # --- Sidebar with Quick Actions ---
    with st.sidebar:
        st.markdown("---")
        st.subheader("🚀 Quick Help")
        
        if st.button("🆕 New Team Member Guide"):
            question = "I'm new to the team, what should I know for my first week?"
            st.session_state.chat_messages.append({"role": "user", "content": question})
            
            with st.spinner("Getting guidance..."):
                response_data = get_ai_response(question, selected_team_id)
            
            if "error" not in response_data:
                st.session_state.chat_messages.append({
                    "role": "assistant", 
                    "content": response_data.get("response", ""),
                    "query_id": response_data.get("query_id"),
                    "suggested_followups": response_data.get("suggested_followups", []),
                    "related_docs": response_data.get("related_docs", [])
                })
            st.rerun()
        
        if st.button("📋 Test Case Standards"):
            question = "What's our standard format for writing test cases?"
            st.session_state.chat_messages.append({"role": "user", "content": question})
            
            with st.spinner("Looking up standards..."):
                response_data = get_ai_response(question, selected_team_id)
            
            if "error" not in response_data:
                st.session_state.chat_messages.append({
                    "role": "assistant", 
                    "content": response_data.get("response", ""),
                    "query_id": response_data.get("query_id"),
                    "suggested_followups": response_data.get("suggested_followups", []),
                    "related_docs": response_data.get("related_docs", [])
                })
            st.rerun()
        
        if st.button("🐛 Bug Classification"):
            question = "How do I classify bug severity and priority?"
            st.session_state.chat_messages.append({"role": "user", "content": question})
            
            with st.spinner("Getting classification guide..."):
                response_data = get_ai_response(question, selected_team_id)
            
            if "error" not in response_data:
                st.session_state.chat_messages.append({
                    "role": "assistant", 
                    "content": response_data.get("response", ""),
                    "query_id": response_data.get("query_id"),
                    "suggested_followups": response_data.get("suggested_followups", []),
                    "related_docs": response_data.get("related_docs", [])
                })
            st.rerun()
        
        if st.button("🔧 Common Issues"):
            question = f"What are common testing issues for {project_name} team?"
            st.session_state.chat_messages.append({"role": "user", "content": question})
            
            with st.spinner("Finding common issues..."):
                response_data = get_ai_response(question, selected_team_id)
            
            if "error" not in response_data:
                st.session_state.chat_messages.append({
                    "role": "assistant", 
                    "content": response_data.get("response", ""),
                    "query_id": response_data.get("query_id"),
                    "suggested_followups": response_data.get("suggested_followups", []),
                    "related_docs": response_data.get("related_docs", [])
                })
            st.rerun()
        
        if st.button("🔄 Clear Chat"):
            st.session_state.chat_messages = []
            st.session_state.show_feedback = {}
            st.rerun()
    
    # --- Main Chat Interface ---
    st.markdown("---")
    
    # Display chat messages
    for i, message in enumerate(st.session_state.chat_messages):
        if message["role"] == "user":
            with st.container():
                st.markdown(f"**👤 You:** {message['content']}")
        
        elif message["role"] == "assistant":
            with st.container():
                st.markdown(f"**🤖 PW QA Assistant:** {message['content']}")
                
                # Show suggested follow-ups
                if message.get("suggested_followups"):
                    st.markdown("**💡 Related Questions:**")
                    col1, col2, col3 = st.columns(3)
                    for idx, followup in enumerate(message["suggested_followups"][:3]):
                        with [col1, col2, col3][idx % 3]:
                            if st.button(followup, key=f"followup_{i}_{idx}"):
                                st.session_state.chat_messages.append({"role": "user", "content": followup})
                                
                                with st.spinner("Getting answer..."):
                                    response_data = get_ai_response(followup, selected_team_id)
                                
                                if "error" not in response_data:
                                    st.session_state.chat_messages.append({
                                        "role": "assistant", 
                                        "content": response_data.get("response", ""),
                                        "query_id": response_data.get("query_id"),
                                        "suggested_followups": response_data.get("suggested_followups", []),
                                        "related_docs": response_data.get("related_docs", [])
                                    })
                                st.rerun()
                
                # Show related documents
                if message.get("related_docs"):
                    with st.expander("📚 Related Documentation"):
                        for doc in message["related_docs"]:
                            st.markdown(f"**{doc.get('title', 'Document')}**")
                            st.write(doc.get('content', '')[:200] + "..." if len(doc.get('content', '')) > 200 else doc.get('content', ''))
                
                # Feedback buttons
                if message.get("query_id") and message["query_id"] not in st.session_state.show_feedback:
                    st.markdown("**Was this helpful?**")
                    col1, col2, col3, col4, col5 = st.columns(5)
                    
                    with col1:
                        if st.button("⭐", key=f"rating_1_{i}", help="Poor"):
                            submit_feedback(message["query_id"], 1)
                            st.session_state.show_feedback[message["query_id"]] = 1
                            st.rerun()
                    
                    with col2:
                        if st.button("⭐⭐", key=f"rating_2_{i}", help="Fair"):
                            submit_feedback(message["query_id"], 2)
                            st.session_state.show_feedback[message["query_id"]] = 2
                            st.rerun()
                    
                    with col3:
                        if st.button("⭐⭐⭐", key=f"rating_3_{i}", help="Good"):
                            submit_feedback(message["query_id"], 3)
                            st.session_state.show_feedback[message["query_id"]] = 3
                            st.rerun()
                    
                    with col4:
                        if st.button("⭐⭐⭐⭐", key=f"rating_4_{i}", help="Very Good"):
                            submit_feedback(message["query_id"], 4)
                            st.session_state.show_feedback[message["query_id"]] = 4
                            st.rerun()
                    
                    with col5:
                        if st.button("⭐⭐⭐⭐⭐", key=f"rating_5_{i}", help="Excellent"):
                            submit_feedback(message["query_id"], 5)
                            st.session_state.show_feedback[message["query_id"]] = 5
                            st.rerun()
                
                elif message.get("query_id") in st.session_state.show_feedback:
                    rating = st.session_state.show_feedback[message["query_id"]]
                    st.markdown(f"✅ **Feedback submitted:** {'⭐' * rating}")
        
        st.markdown("---")
    
    # --- Chat Input ---
    if selected_team_id:
        user_input = st.text_input(
            "Ask me anything about QA processes, testing procedures, or team-specific guidance...",
            key="chat_input",
            placeholder=f"e.g., How do I test user stories for {project_name}?"
        )
        
        col1, col2 = st.columns([6, 1])
        
        with col1:
            if st.button("📤 Send", key="send_message") and user_input.strip():
                # Add user message
                st.session_state.chat_messages.append({"role": "user", "content": user_input})
                
                # Get AI response
                with st.spinner("🤖 Thinking..."):
                    response_data = get_ai_response(user_input, selected_team_id)
                
                # Add assistant response
                if "error" in response_data:
                    st.session_state.chat_messages.append({
                        "role": "assistant", 
                        "content": f"⚠️ Sorry, I encountered an issue: {response_data['error']}. Please try again or contact your team lead for help."
                    })
                else:
                    st.session_state.chat_messages.append({
                        "role": "assistant", 
                        "content": response_data.get("response", ""),
                        "query_id": response_data.get("query_id"),
                        "suggested_followups": response_data.get("suggested_followups", []),
                        "related_docs": response_data.get("related_docs", [])
                    })
                
                st.rerun()
        
        with col2:
            if st.button("🎯 Examples"):
                examples = [
                    "How do I write effective test cases?",
                    "What's our bug reporting process?",
                    "How do I access test environments?",
                    "What should I test for API endpoints?",
                    "How do I handle test data setup?"
                ]
                
                st.markdown("**Example Questions:**")
                for example in examples:
                    if st.button(example, key=f"example_{hash(example)}"):
                        st.session_state.chat_messages.append({"role": "user", "content": example})
                        
                        with st.spinner("Getting answer..."):
                            response_data = get_ai_response(example, selected_team_id)
                        
                        if "error" not in response_data:
                            st.session_state.chat_messages.append({
                                "role": "assistant", 
                                "content": response_data.get("response", ""),
                                "query_id": response_data.get("query_id"),
                                "suggested_followups": response_data.get("suggested_followups", []),
                                "related_docs": response_data.get("related_docs", [])
                            })
                        st.rerun()
    else:
        st.warning("⚠️ Please select a team from the sidebar to start chatting.")

# --- Footer ---
st.markdown(
    """
    <hr style="margin-top: 50px; margin-bottom: 10px;">
    <div style="text-align: center; color: grey; font-size: 14px;">
        © 2025 PW QA Assistant Platform. Built with ❤️ by Team 3
    </div>
    """,
    unsafe_allow_html=True
)